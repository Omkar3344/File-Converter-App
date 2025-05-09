import os
import time
from PIL import Image
import pandas as pd
import tempfile
import io
import sys
import subprocess
from PyPDF2 import PdfReader, PdfWriter
from docx import Document
import traceback
import logging
import json

# Set up logging
logging.basicConfig(level=logging.INFO, 
                    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
                    handlers=[logging.FileHandler("converter.log"), logging.StreamHandler()])
logger = logging.getLogger("FileConverter")

# Map of supported input formats and their possible output formats
FORMAT_MAP = {
    # Image formats
    '.jpg': ['png', 'webp', 'bmp', 'gif', 'pdf'],
    '.jpeg': ['png', 'webp', 'bmp', 'gif', 'pdf'],
    '.png': ['jpg', 'webp', 'bmp', 'gif', 'pdf'],
    '.bmp': ['jpg', 'png', 'webp', 'gif', 'pdf'],
    '.webp': ['jpg', 'png', 'bmp', 'gif', 'pdf'],
    '.gif': ['jpg', 'png', 'webp', 'bmp', 'pdf'],
    
    # Document formats
    '.pdf': ['txt', 'jpg', 'png'],
    '.docx': ['pdf', 'txt'],
    '.txt': ['pdf', 'docx'],
    
    # Data formats
    '.csv': ['xlsx', 'json', 'xml', 'html'],
    '.xlsx': ['csv', 'json', 'xml', 'html'],
    '.json': ['csv', 'xlsx', 'xml', 'html'],
}

def get_available_formats(file_extension):
    """Get available output formats for a given file extension."""
    return FORMAT_MAP.get(file_extension.lower(), [])

def convert_file(source_path, target_path, progress_callback=None):
    """
    Convert a file from one format to another.
    
    Args:
        source_path: Path to the source file
        target_path: Path where the converted file should be saved
        progress_callback: Function to call with progress updates (0-100)
    
    Returns:
        bool: True if conversion was successful, False otherwise
    """
    try:
        # Validate file paths
        if not os.path.exists(source_path):
            logger.error(f"Source file does not exist: {source_path}")
            return False
            
        # Create target directory if it doesn't exist
        target_dir = os.path.dirname(target_path)
        if target_dir and not os.path.exists(target_dir):
            os.makedirs(target_dir)
            
        source_ext = os.path.splitext(source_path)[1].lower()
        target_ext = os.path.splitext(target_path)[1].lower()
        
        if target_ext.startswith('.'):
            target_ext = target_ext[1:]
            
        logger.info(f"Converting {source_path} ({source_ext}) to {target_path} ({target_ext})")
        
        # Simulate progress updates
        def update_progress(value):
            if progress_callback:
                progress_callback(value)
                
        update_progress(10)
        
        # Handle image conversions
        if source_ext.lower() in ['.jpg', '.jpeg', '.png', '.bmp', '.webp', '.gif']:
            if target_ext.lower() in ['jpg', 'jpeg', 'png', 'bmp', 'webp', 'gif']:
                return convert_image(source_path, target_path, update_progress)
            elif target_ext.lower() == 'pdf':
                return image_to_pdf(source_path, target_path, update_progress)
        
        # Handle document conversions
        elif source_ext.lower() == '.pdf':
            if target_ext.lower() in ['jpg', 'jpeg', 'png']:
                return pdf_to_images(source_path, target_path, update_progress)
            elif target_ext.lower() == 'txt':
                return pdf_to_text(source_path, target_path, update_progress)
        
        elif source_ext.lower() == '.docx':
            if target_ext.lower() == 'pdf':
                return docx_to_pdf(source_path, target_path, update_progress)
            elif target_ext.lower() == 'txt':
                return docx_to_text(source_path, target_path, update_progress)
        
        elif source_ext.lower() == '.txt':
            if target_ext.lower() == 'pdf':
                return text_to_pdf(source_path, target_path, update_progress)
            elif target_ext.lower() == 'docx':
                return text_to_docx(source_path, target_path, update_progress)
        
        # Handle data formats
        elif source_ext.lower() in ['.csv', '.xlsx', '.json']:
            if target_ext.lower() in ['csv', 'xlsx', 'json', 'xml', 'html']:
                return convert_data_format(source_path, target_ext, target_path, update_progress)
        
        logger.warning(f"No specific conversion routine found for {source_ext} to {target_ext}")
        update_progress(100)
        return False
    
    except Exception as e:
        logger.error(f"Conversion error: {str(e)}")
        logger.error(traceback.format_exc())
        return False

def convert_image(source_path, target_path, progress_callback):
    """Convert between image formats."""
    try:
        progress_callback(20)
        
        # Open the image file
        img = Image.open(source_path)
        logger.info(f"Image opened: {source_path}, Mode: {img.mode}, Size: {img.size}")
        
        # Handle special cases for different formats
        target_ext = os.path.splitext(target_path)[1].lower()
        
        # JPEG conversion - must be RGB
        if target_ext in ['.jpg', '.jpeg'] or target_path.lower().endswith(('.jpg', '.jpeg')):
            logger.info("Converting to JPEG format (ensuring RGB mode)")
            # Convert to RGB mode for JPEG (which doesn't support alpha channels or palettes)
            if img.mode != 'RGB':
                img = img.convert('RGB')
                logger.info(f"Converted image mode to RGB")
        
        # PNG conversion - can handle RGBA
        elif target_ext == '.png' or target_path.lower().endswith('.png'):
            # PNG supports various modes, but RGBA is most common for transparency
            if img.mode == 'P' and 'transparency' in img.info:
                img = img.convert('RGBA')
                logger.info(f"Converted palette with transparency to RGBA")
        
        # GIF conversion
        elif target_ext == '.gif' or target_path.lower().endswith('.gif'):
            # For animated GIFs, we'd need special handling here
            # For single frame, convert to P mode with adaptive palette
            if img.mode not in ['P', 'L', 'RGB']:
                img = img.convert('RGB').convert('P', palette=Image.ADAPTIVE)
                logger.info(f"Converted to palette mode for GIF")
        
        progress_callback(60)
        
        # Save the converted image
        logger.info(f"Saving image to {target_path}")
        img.save(target_path)
        
        # Verify the file was created
        if not os.path.exists(target_path):
            logger.error(f"Failed to save image: File not created at {target_path}")
            return False
            
        file_size = os.path.getsize(target_path)
        if file_size == 0:
            logger.error(f"Saved file has zero size: {target_path}")
            return False
            
        logger.info(f"Successfully saved image: {target_path}, Size: {file_size} bytes")
        progress_callback(100)
        return True
        
    except Exception as e:
        logger.error(f"Image conversion error: {str(e)}")
        logger.error(traceback.format_exc())
        return False

def image_to_pdf(source_path, target_path, progress_callback):
    """Convert an image to PDF."""
    try:
        progress_callback(20)
        
        img = Image.open(source_path)
        logger.info(f"Converting image to PDF: {source_path} -> {target_path}")
        
        # Always convert to RGB for PDF
        if img.mode != 'RGB':
            img = img.convert('RGB')
            logger.info("Converted image to RGB mode for PDF")
        
        progress_callback(60)
        
        # Save as PDF
        img.save(target_path, "PDF", resolution=100.0)
        
        # Verify file creation
        if not os.path.exists(target_path):
            logger.error(f"Failed to create PDF: {target_path}")
            return False
            
        logger.info(f"Successfully created PDF: {target_path}")
        progress_callback(100)
        return True
        
    except Exception as e:
        logger.error(f"Image to PDF conversion error: {str(e)}")
        logger.error(traceback.format_exc())
        return False

def pdf_to_images(source_path, target_path, progress_callback):
    """Extract pages from a PDF as images."""
    try:
        progress_callback(10)
        
        logger.info(f"Converting PDF to image(s): {source_path} -> {target_path}")
        
        # Check if pdf2image is available
        try:
            from pdf2image import convert_from_path
            has_pdf2image = True
        except ImportError:
            has_pdf2image = False
            logger.warning("pdf2image library not available, falling back to PyMuPDF")
        
        # Check if PyMuPDF is available
        try:
            import fitz
            has_pymupdf = True
        except ImportError:
            has_pymupdf = False
            if not has_pdf2image:
                logger.error("Neither pdf2image nor PyMuPDF is available")
                return False
        
        # Get base name without extension
        base_path = os.path.splitext(target_path)[0]
        ext = os.path.splitext(target_path)[1][1:]  # Remove leading dot
        
        progress_callback(20)
        
        # Convert using pdf2image if available
        if has_pdf2image:
            try:
                images = convert_from_path(source_path, dpi=300)
                total_pages = len(images)
                logger.info(f"PDF has {total_pages} pages")
                
                for i, img in enumerate(images):
                    img_path = f"{base_path}_page{i+1}.{ext}"
                    img.save(img_path)
                    logger.info(f"Created image for page {i+1}: {img_path}")
                    progress_callback(20 + (i+1) * 70 // total_pages)
                    
                return True
            except Exception as e:
                logger.error(f"pdf2image conversion failed: {str(e)}")
                if not has_pymupdf:
                    return False
        
        # Convert using PyMuPDF if pdf2image failed or is not available
        if has_pymupdf:
            pdf_document = fitz.open(source_path)
            total_pages = len(pdf_document)
            logger.info(f"PDF has {total_pages} pages")
            
            for i, page in enumerate(pdf_document):
                pix = page.get_pixmap(matrix=fitz.Matrix(300/72, 300/72))
                img_path = f"{base_path}_page{i+1}.{ext}"
                
                # For JPG, we need RGB
                if ext.lower() in ["jpg", "jpeg"]:
                    pix = fitz.Pixmap(fitz.csRGB, pix)
                
                pix.save(img_path)
                logger.info(f"Created image for page {i+1}: {img_path}")
                progress_callback(20 + (i+1) * 70 // total_pages)
            
            pdf_document.close()
            return True
            
        # If we got here, both methods failed
        logger.error("Failed to convert PDF to images")
        return False
        
    except Exception as e:
        logger.error(f"PDF to images error: {str(e)}")
        logger.error(traceback.format_exc())
        return False

def pdf_to_text(source_path, target_path, progress_callback):
    """Extract text from a PDF using a combination of methods for better results."""
    try:
        progress_callback(10)
        
        logger.info(f"Converting PDF to text: {source_path} -> {target_path}")
        
        # Try PyPDF2 first
        pdf = PdfReader(source_path)
        total_pages = len(pdf.pages)
        logger.info(f"PDF has {total_pages} pages")
        
        # Check if PyMuPDF is available for better text extraction
        try:
            import fitz
            has_pymupdf = True
        except ImportError:
            has_pymupdf = False
            logger.warning("PyMuPDF not available, using PyPDF2 for text extraction")
        
        progress_callback(20)
        
        with open(target_path, 'w', encoding='utf-8') as text_file:
            # Use PyMuPDF if available
            if has_pymupdf:
                doc = fitz.open(source_path)
                for i, page in enumerate(doc):
                    text = page.get_text()
                    text_file.write(text)
                    text_file.write('\n\n--- Page Break ---\n\n')
                    logger.info(f"Extracted text from page {i+1} (length: {len(text)} chars)")
                    progress_callback(20 + (i+1) * 70 // total_pages)
                doc.close()
            else:
                # Fallback to PyPDF2
                for i, page in enumerate(pdf.pages):
                    text = page.extract_text()
                    text_file.write(text)
                    text_file.write('\n\n--- Page Break ---\n\n')
                    logger.info(f"Extracted text from page {i+1} (length: {len(text)} chars)")
                    progress_callback(20 + (i+1) * 70 // total_pages)
        
        # Verify file was created
        if not os.path.exists(target_path):
            logger.error(f"Failed to create text file: {target_path}")
            return False
            
        # If file is empty, try alternative method with pdfminer.six if available
        if os.path.getsize(target_path) < 100:  # If file is very small
            logger.warning("Text extraction yielded little content, trying alternative method")
            try:
                from pdfminer.high_level import extract_text as pdfminer_extract_text
                text = pdfminer_extract_text(source_path)
                with open(target_path, 'w', encoding='utf-8') as text_file:
                    text_file.write(text)
                logger.info(f"Used pdfminer.six to extract text (length: {len(text)} chars)")
            except ImportError:
                logger.warning("pdfminer.six not available for alternative text extraction")
            
        logger.info(f"Successfully created text file: {target_path}")
        progress_callback(100)
        return True
        
    except Exception as e:
        logger.error(f"PDF to text error: {str(e)}")
        logger.error(traceback.format_exc())
        return False

def docx_to_pdf(source_path, target_path, progress_callback):
    """Convert DOCX to PDF using multiple methods."""
    try:
        progress_callback(20)
        
        logger.info(f"Converting DOCX to PDF: {source_path} -> {target_path}")
        
        # Try docx2pdf if available
        try:
            from docx2pdf import convert
            logger.info("Using docx2pdf for conversion")
            convert(source_path, target_path)
            
            # Verify file was created
            if os.path.exists(target_path) and os.path.getsize(target_path) > 0:
                logger.info(f"Successfully created PDF file using docx2pdf: {target_path}")
                progress_callback(100)
                return True
            else:
                logger.warning("docx2pdf conversion failed or created empty file")
        except ImportError:
            logger.warning("docx2pdf not available, trying alternative method")
        
        # Alternative method using python-docx and reportlab
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.platypus import SimpleDocTemplate, Paragraph
            
            logger.info("Using python-docx and reportlab for conversion")
            
            # Read the DOCX file
            doc = Document(source_path)
            
            # Extract text content
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            
            # Create a PDF
            styles = getSampleStyleSheet()
            pdf_doc = SimpleDocTemplate(target_path, pagesize=letter)
            
            # Convert paragraphs to reportlab paragraphs
            pdf_elements = [Paragraph(text, styles["Normal"]) for text in paragraphs]
            
            # Build the PDF
            pdf_doc.build(pdf_elements)
            
            # Verify file was created
            if os.path.exists(target_path) and os.path.getsize(target_path) > 0:
                logger.info(f"Successfully created PDF file using reportlab: {target_path}")
                progress_callback(100)
                return True
            else:
                logger.warning("reportlab conversion failed or created empty file")
        except ImportError:
            logger.warning("reportlab not available, falling back to basic method")
        
        # Last resort - create a basic PDF with text
        progress_callback(80)
        
        # Read DOCX content
        doc = Document(source_path)
        content = "\n".join([para.text for para in doc.paragraphs])
        
        # Create a simple PDF with PyPDF2
        from PyPDF2 import PdfWriter
        from reportlab.pdfgen import canvas
        from io import BytesIO
        
        # Create a PDF in memory
        packet = BytesIO()
        can = canvas.Canvas(packet, pagesize=letter)
        
        # Write the text - basic implementation
        text_lines = content.split('\n')
        y = 750  # starting y position
        for line in text_lines:
            if not line.strip():
                y -= 12
                continue
                
            # Split long lines
            words = line.split()
            current_line = ""
            for word in words:
                test_line = current_line + " " + word if current_line else word
                if can.stringWidth(test_line, "Helvetica", 11) < 500:  # max width
                    current_line = test_line
                else:
                    can.drawString(50, y, current_line)
                    y -= 12
                    if y < 50:  # new page
                        can.showPage()
                        y = 750
                    current_line = word
            
            if current_line:
                can.drawString(50, y, current_line)
                y -= 12
            
            if y < 50:  # new page
                can.showPage()
                y = 750
        
        can.save()
        
        # Get the PDF data and write to file
        packet.seek(0)
        new_pdf = PdfReader(packet)
        output = PdfWriter()
        
        # Add all pages
        for page in range(len(new_pdf.pages)):
            output.add_page(new_pdf.pages[page])
        
        # Write to file
        with open(target_path, "wb") as output_file:
            output.write(output_file)
        
        # Verify file was created
        if not os.path.exists(target_path):
            logger.error(f"Failed to create PDF file: {target_path}")
            return False
            
        logger.info(f"Successfully created basic PDF file: {target_path}")
        progress_callback(100)
        return True
        
    except Exception as e:
        logger.error(f"DOCX to PDF error: {str(e)}")
        logger.error(traceback.format_exc())
        return False

def docx_to_text(source_path, target_path, progress_callback):
    """Extract text from a DOCX file."""
    try:
        progress_callback(20)
        
        logger.info(f"Converting DOCX to text: {source_path} -> {target_path}")
        doc = Document(source_path)
        
        progress_callback(50)
        
        with open(target_path, 'w', encoding='utf-8') as text_file:
            # Extract text from paragraphs
            for para in doc.paragraphs:
                text_file.write(para.text)
                text_file.write('\n')
            
            # Extract text from tables
            for table in doc.tables:
                text_file.write('\n--- TABLE ---\n')
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    text_file.write(' | '.join(row_text))
                    text_file.write('\n')
                text_file.write('--- END TABLE ---\n\n')
        
        # Verify file was created
        if not os.path.exists(target_path):
            logger.error(f"Failed to create text file: {target_path}")
            return False
            
        logger.info(f"Successfully created text file: {target_path}")
        progress_callback(100)
        return True
        
    except Exception as e:
        logger.error(f"DOCX to text error: {str(e)}")
        logger.error(traceback.format_exc())
        return False

def text_to_pdf(source_path, target_path, progress_callback):
    """Convert plain text to PDF."""
    try:
        progress_callback(20)
        
        logger.info(f"Converting text to PDF: {source_path} -> {target_path}")
        
        # Try using reportlab if available
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.platypus import SimpleDocTemplate, Paragraph
            
            logger.info("Using reportlab for text to PDF conversion")
            
            # Read the text file
            with open(source_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Split into paragraphs
            paragraphs = content.split('\n')
            
            # Create a PDF
            styles = getSampleStyleSheet()
            pdf_doc = SimpleDocTemplate(target_path, pagesize=letter)
            
            # Convert paragraphs to reportlab paragraphs
            pdf_elements = []
            for text in paragraphs:
                if text.strip():  # Skip empty lines
                    pdf_elements.append(Paragraph(text, styles["Normal"]))
                else:
                    # Add spacer for empty lines
                    pdf_elements.append(Paragraph("&nbsp;", styles["Normal"]))
            
            # Build the PDF
            pdf_doc.build(pdf_elements)
            
            # Verify file was created
            if os.path.exists(target_path) and os.path.getsize(target_path) > 0:
                logger.info(f"Successfully created PDF file using reportlab: {target_path}")
                progress_callback(100)
                return True
            else:
                logger.warning("reportlab conversion failed or created empty file")
            
        except ImportError:
            logger.warning("reportlab not available, using basic method")
        
        # Basic method using PyPDF2 and canvas
        with open(source_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        progress_callback(50)
        
        # Create a PDF using PyPDF2
        from PyPDF2 import PdfWriter
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        from io import BytesIO
        
        # Create a PDF in memory
        packet = BytesIO()
        can = canvas.Canvas(packet, pagesize=letter)
        
        # Write the text - basic implementation
        text_lines = content.split('\n')
        y = 750  # starting y position
        for line in text_lines:
            if not line.strip():
                y -= 12
                continue
                
            # Split long lines
            words = line.split()
            current_line = ""
            for word in words:
                test_line = current_line + " " + word if current_line else word
                if can.stringWidth(test_line, "Helvetica", 11) < 500:  # max width
                    current_line = test_line
                else:
                    can.drawString(50, y, current_line)
                    y -= 12
                    if y < 50:  # new page
                        can.showPage()
                        y = 750
                    current_line = word
            
            if current_line:
                can.drawString(50, y, current_line)
                y -= 12
            
            if y < 50:  # new page
                can.showPage()
                y = 750
        
        can.save()
        
        # Get the PDF data and write to file
        packet.seek(0)
        new_pdf = PdfReader(packet)
        output = PdfWriter()
        
        # Add all pages
        for page in range(len(new_pdf.pages)):
            output.add_page(new_pdf.pages[page])
        
        # Write to file
        with open(target_path, "wb") as output_file:
            output.write(output_file)
        
        progress_callback(80)
        
        # Verify file was created
        if not os.path.exists(target_path):
            logger.error(f"Failed to create PDF file: {target_path}")
            return False
            
        logger.info(f"Successfully created PDF file: {target_path}")
        progress_callback(100)
        return True
        
    except Exception as e:
        logger.error(f"Text to PDF error: {str(e)}")
        logger.error(traceback.format_exc())
        return False

def text_to_docx(source_path, target_path, progress_callback):
    """Convert plain text to DOCX."""
    try:
        progress_callback(20)
        
        logger.info(f"Converting text to DOCX: {source_path} -> {target_path}")
        with open(source_path, 'r', encoding='utf-8') as text_file:
            text = text_file.read()
        
        progress_callback(50)
        
        doc = Document()
        # Split text into paragraphs by newlines
        paragraphs = text.split('\n')
        
        for paragraph in paragraphs:
            # Skip empty paragraphs but add an empty paragraph for spacing
            if paragraph.strip():
                doc.add_paragraph(paragraph)
            else:
                doc.add_paragraph()
        
        progress_callback(80)
        
        doc.save(target_path)
        
        # Verify file was created
        if not os.path.exists(target_path):
            logger.error(f"Failed to create DOCX file: {target_path}")
            return False
            
        logger.info(f"Successfully created DOCX file: {target_path}")
        progress_callback(100)
        return True
        
    except Exception as e:
        logger.error(f"Text to DOCX error: {str(e)}")
        logger.error(traceback.format_exc())
        return False

def convert_data_format(source_path, target_format, target_path, progress_callback):
    """Convert between data formats (CSV, Excel, JSON, etc.)."""
    try:
        progress_callback(10)
        
        logger.info(f"Converting data format: {source_path} -> {target_path} ({target_format})")
        # Read source file based on its extension
        source_ext = os.path.splitext(source_path)[1].lower()
        
        if source_ext == '.csv':
            logger.info("Reading CSV file")
            df = pd.read_csv(source_path)
        elif source_ext == '.xlsx':
            logger.info("Reading Excel file")
            df = pd.read_excel(source_path)
        elif source_ext == '.json':
            logger.info("Reading JSON file")
            # Try to handle different JSON formats
            try:
                df = pd.read_json(source_path)
            except ValueError:
                # For JSON Lines or irregular formats
                with open(source_path, 'r', encoding='utf-8') as f:
                    json_data = json.load(f)
                
                # Handle different JSON structures
                if isinstance(json_data, list):
                    df = pd.DataFrame(json_data)
                elif isinstance(json_data, dict):
                    # If it's a nested dictionary, try to normalize it
                    if any(isinstance(v, dict) for v in json_data.values()):
                        df = pd.json_normalize(json_data)
                    else:
                        df = pd.DataFrame([json_data])
                else:
                    logger.error("Unsupported JSON structure")
                    return False
        else:
            logger.error(f"Unsupported source format: {source_ext}")
            return False
        
        logger.info(f"Read data with shape: {df.shape}")
        progress_callback(50)
        
        # Write to target format
        if target_format == 'csv':
            logger.info("Writing to CSV format")
            df.to_csv(target_path, index=False)
        elif target_format == 'xlsx':
            logger.info("Writing to Excel format")
            df.to_excel(target_path, index=False)
        elif target_format == 'json':
            logger.info("Writing to JSON format")
            df.to_json(target_path, orient='records')
        elif target_format == 'xml':
            logger.info("Writing to XML format")
            # Check if to_xml is available (pandas >= 1.3.0)
            if hasattr(df, 'to_xml'):
                df.to_xml(target_path)
            else:
                # Fallback for older pandas versions
                xml_data = '<?xml version="1.0" encoding="UTF-8"?>\n<data>\n'
                for _, row in df.iterrows():
                    xml_data += '  <record>\n'
                    for col, value in row.items():
                        xml_data += f'    <{col}>{value}</{col}>\n'
                    xml_data += '  </record>\n'
                xml_data += '</data>'
                
                with open(target_path, 'w', encoding='utf-8') as f:
                    f.write(xml_data)
        elif target_format == 'html':
            logger.info("Writing to HTML format")
            df.to_html(target_path, index=False)
        else:
            logger.error(f"Unsupported target format: {target_format}")
            return False
        
        # Verify file was created
        if not os.path.exists(target_path):
            logger.error(f"Failed to create file: {target_path}")
            return False
            
        logger.info(f"Successfully created {target_format} file: {target_path}")
        progress_callback(100)
        return True
        
    except Exception as e:
        logger.error(f"Data format conversion error: {str(e)}")
        logger.error(traceback.format_exc())
        return False